"""Рендер в DOCX (python-docx) по ГОСТ 2.105."""

from __future__ import annotations

from pathlib import Path

from ..parser.model import Project
from ..styles.gost_styles import STYLE
from ..styles.docs import get_doc
from .bodies import body_for


def _add_page_numbers(doc) -> None:
    """Номер страницы по центру нижнего поля (ГОСТ 2.105)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = 1  # CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar"); fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar"); fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def _section_num(i: int) -> str:
    return f"{i}"


def render_docx(proj: Project, out_path: Path, diagrams: dict | None = None,
                doc_code: str = "19.402") -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Mm, Pt, Cm

    doc_spec = get_doc(doc_code)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin = Mm(STYLE.margin_left_mm)
    sec.right_margin = Mm(STYLE.margin_right_mm)
    sec.top_margin = Mm(STYLE.margin_top_mm)
    sec.bottom_margin = Mm(STYLE.margin_bottom_mm)

    _add_page_numbers(doc)

    style = doc.styles["Normal"]
    style.font.name = STYLE.font_ru
    style.font.size = Pt(STYLE.font_size)
    style.paragraph_format.line_spacing = STYLE.line_spacing
    style.paragraph_format.first_line_indent = Cm(STYLE.first_line_indent_cm)

    # --- титульный лист ---
    for _ in range(4):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(proj.organisation.upper()).bold = True
    for _ in range(4):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(proj.name)
    r.font.size = Pt(18)
    r.bold = True
    doc.add_paragraph(doc_spec.title)
    doc.add_paragraph(doc_spec.subtitle)
    for _ in range(5):
        doc.add_paragraph("")
    doc.add_paragraph(f"Разработал: {proj.author or '________'}")
    doc.add_paragraph(f"Дата: {__import__('datetime').date.today().isoformat()}")
    doc.add_paragraph(f"Документ: ГОСТ {doc_spec.code} (ЕСПД)")
    doc.add_page_break()

    # --- аннотация ---
    h = doc.add_paragraph()
    h.add_run("Аннотация").bold = True
    doc.add_paragraph(
        f"Настоящий документ — «{doc_spec.subtitle}» для программы «{proj.name}»; "
        f"соответствует требованиям ЕСПД (ГОСТ 19) и ГОСТ 2.105. "
        f"{proj.comment}".strip())

    # --- содержание (краткое) ---
    h = doc.add_paragraph()
    h.add_run("Содержание").bold = True
    for num, title, _key in doc_spec.sections:
        doc.add_paragraph(f"{num}. {title}")

    # --- разделы ---
    for num, title, key in doc_spec.sections:
        h = doc.add_paragraph()
        h.add_run(f"{num}. {title}").bold = True
        doc.add_paragraph(body_for(proj, key))

    doc.add_page_break()

    # --- приложение: структура кода ---
    h = doc.add_paragraph()
    h.add_run("Приложение А. Структура программы").bold = True

    b = proj.build
    if b.kind:
        p = doc.add_paragraph()
        p.add_run(f"Система сборки: {'CMake' if b.kind == 'cmake' else 'qmake'}"
                  + (f" {b.version}" if b.version else "")).bold = True
        if b.targets:
            doc.add_paragraph("Цели: " + ", ".join(b.targets))
        if b.qt_modules:
            doc.add_paragraph("Модули Qt: " + ", ".join(b.qt_modules))
        if b.dependencies:
            doc.add_paragraph("Зависимости: " + ", ".join(b.dependencies))

    if diagrams and diagrams.get("classes"):
        doc.add_picture(str(diagrams["classes"]), width=Cm(16))
        doc.add_paragraph("Рисунок А.1 — Диаграмма классов")
    if diagrams and diagrams.get("calls"):
        doc.add_picture(str(diagrams["calls"]), width=Cm(16))
        doc.add_paragraph("Рисунок А.2 — Граф вызовов")
    flows = (diagrams or {}).get("flows") or {}
    fig = 3
    for fname in sorted(flows):
        doc.add_picture(str(flows[fname]), width=Cm(12))
        doc.add_paragraph(f"Рисунок А.{fig} — Блок-схема функции {fname}")
        fig += 1
    seqs = (diagrams or {}).get("seq") or {}
    for entry in sorted(seqs):
        doc.add_picture(str(seqs[entry]), width=Cm(15))
        doc.add_paragraph(f"Рисунок А.{fig} — Диаграмма последовательности (вход: {entry})")
        fig += 1

    for cls in proj.classes:
        doc.add_paragraph(f"Класс {cls.name}" + (f" : {cls.base}" if cls.base else "")).bold = True
        if cls.comment:
            doc.add_paragraph(cls.comment)
        if cls.fields:
            doc.add_paragraph("Поля:")
            for f_ in cls.fields:
                doc.add_paragraph(f_ + " —", style="List Bullet")
        if cls.methods:
            doc.add_paragraph("Методы:")
            for m in cls.methods:
                kind = {"signal": "сигнал", "slot": "слот"}.get(m.kind, "метод")
                doc.add_paragraph(
                    f"{kind}: {m.return_type} {m.name}({', '.join(m.params)})",
                    style="List Bullet")
                if m.brief:
                    p = doc.add_paragraph(m.brief, style="List Bullet 2")
                    for run in p.runs:
                        run.italic = True

    for fn in proj.functions:
        doc.add_paragraph(f"Функция {fn.name}").bold = True
        doc.add_paragraph(f"{fn.return_type} {fn.name}({', '.join(fn.params)}) — файл {fn.file}:{fn.line}")

    if proj.nn_results:
        h = doc.add_paragraph()
        h.add_run("Приложение Б. Результаты нейросетевого анализа (23 сети)").bold = True
        current = None
        for r in proj.nn_results:
            if r.category != current:
                current = r.category
                p = doc.add_paragraph()
                p.add_run(r.category.upper()).bold = True
            p = doc.add_paragraph()
            p.add_run(f"{r.net_name} [score={r.score}]").bold = True
            doc.add_paragraph(r.text, style="List Bullet")

    if doc_code == "19.401" and proj.sources:
        h = doc.add_paragraph()
        h.add_run("Приложение В. Текст программы").bold = True
        from docx.shared import Pt as _Pt
        for src in proj.sources:
            p = doc.add_paragraph()
            p.add_run(f"Файл: {src}").bold = True
            try:
                code = Path(src).read_text(encoding="utf-8", errors="ignore")
            except OSError:
                continue
            for i, line in enumerate(code.splitlines(), start=1):
                lp = doc.add_paragraph(f"{i:>5} | {line}")
                lp.paragraph_format.first_line_indent = Cm(0)
                lp.paragraph_format.line_spacing = 1.0
                for run in lp.runs:
                    run.font.name = "Courier New"
                    run.font.size = _Pt(10)

    from .bodies import REV_HEADERS, revision_rows
    doc.add_page_break()
    h = doc.add_paragraph()
    h.add_run("Лист регистрации изменений").bold = True
    table = doc.add_table(rows=1, cols=len(REV_HEADERS))
    table.style = "Table Grid"
    for j, htext in enumerate(REV_HEADERS):
        table.rows[0].cells[j].text = htext
    for row in revision_rows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = val

    doc.save(str(out_path))
    return out_path
